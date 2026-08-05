from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Mau_Bao_Cao_DoAn_LapTrinhPythonNangCao.docx"
BACKUP = ROOT / "Mau_Bao_Cao_DoAn_LapTrinhPythonNangCao.template-backup.docx"


def set_font(run, size=13, bold=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def shade(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def remove_template_body_after_toc(doc: Document):
    marker = next(
        p for p in doc.paragraphs
        if "Update Field" in p.text or "cập nhật tự động" in p.text or "Mục lục tự động" in p.text
    )
    body = doc._element.body
    children = list(body)
    start = children.index(marker._p)
    for child in children[start + 1:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    marker.text = "[Cập nhật Mục lục tự động sau khi hoàn thiện báo cáo.]"
    marker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in marker.runs:
        set_font(run, 11)
        run.italic = True


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)
    for name, size in (("Heading 1", 15), ("Heading 2", 14), ("Heading 3", 13)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True


def paragraph(doc: Document, text="", style=None, align=None, first_line=True):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.8)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("• " + text)
    set_font(r)
    return p


def code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(6)
    for index, line in enumerate(text.splitlines()):
        if index:
            p.add_run("\n")
        r = p.add_run(line)
        set_font(r, size=9, name="Consolas")
    return p


def caption(doc: Document, text: str):
    p = paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    for run in p.runs:
        run.italic = True
        set_font(run, 11)
    return p


def add_table(doc: Document, caption_text: str, headers: list[str], rows: list[list[str]], widths=None):
    caption(doc, caption_text)
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, 10, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                for run in p.runs:
                    set_font(run, 9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    paragraph(doc, "", first_line=False)
    return table


def add_content(doc: Document):
    doc.add_page_break()

    doc.add_heading("1. GIỚI THIỆU ĐỀ TÀI", level=1)
    doc.add_heading("1.1. Mô tả bài toán", level=2)
    paragraph(doc, "Nhiều phòng khám nha khoa nhỏ vẫn tiếp nhận lịch hẹn qua điện thoại, tin nhắn hoặc sổ giấy. Cách làm này gây khó khăn khi tra cứu, dễ nhầm lẫn khung giờ và không hỗ trợ khách hàng ngoài giờ hành chính. Đề tài xây dựng website quản lý phòng khám nha khoa NALI, tập trung số hóa quy trình giới thiệu dịch vụ, quản lý người dùng và đặt lịch hẹn trực tuyến.")
    paragraph(doc, "Hệ thống có ba nhóm người dùng. Khách chưa đăng nhập có thể xem dịch vụ, tìm hiểu đội ngũ và trao đổi với trợ lý AI. Thành viên đã đăng nhập có thể đặt lịch và theo dõi lịch của mình. Quản trị viên có thể quản lý dịch vụ, lịch hẹn, khách hàng, phản hồi và các chỉ số tổng quan.")

    doc.add_heading("1.2. Mục tiêu và phạm vi", level=2)
    bullet(doc, "Xây dựng ứng dụng web Flask có đăng ký/đăng nhập, phân quyền, CRUD dịch vụ, tìm kiếm và phân trang.")
    bullet(doc, "Áp dụng Jinja2, Flask-WTF, Flask-SQLAlchemy, Flask-Mail và lập trình mạng bằng urllib/JSON.")
    bullet(doc, "Tích hợp trợ lý AI có RAG, hỗ trợ tư vấn và đặt lịch theo hội thoại.")
    bullet(doc, "Đóng gói các thành phần web, cơ sở dữ liệu và AI bằng Docker để thuận tiện triển khai.")
    paragraph(doc, "Phạm vi đề tài giới hạn ở nghiệp vụ dịch vụ nha khoa, khách hàng và lịch hẹn. Hệ thống chưa tích hợp thanh toán trực tuyến thật, hồ sơ bệnh án điện tử chuyên sâu hoặc kết nối thiết bị y tế.")

    doc.add_heading("1.3. Công nghệ sử dụng", level=2)
    add_table(doc, "Bảng 1.1. Công nghệ sử dụng", ["Nhóm", "Công nghệ và vai trò"], [
        ["Ngôn ngữ", "Python 3.11; PHP 8.2 cho giao diện tương thích hiện có"],
        ["Web", "Flask 3.1, app factory, Blueprint, Jinja2"],
        ["Biểu mẫu", "Flask-WTF, WTForms, CSRF, validator"],
        ["Dữ liệu", "Flask-SQLAlchemy, PyMySQL, MySQL 8"],
        ["Bảo mật", "Flask-Login, bcrypt, phân quyền admin"],
        ["Mạng và email", "urllib, JSON, REST API, Flask-Mail/SMTP"],
        ["AI", "FastAPI, Qwen2.5-3B, QLoRA, RAG TF-IDF/embedding, Ollama"],
        ["Triển khai", "Docker Compose, Git/GitHub"],
    ], [4.3, 12.7])

    doc.add_heading("1.4. Phân công công việc trong nhóm", level=2)
    add_table(doc, "Bảng 1.2. Phân công công việc", ["Thành viên", "Nhiệm vụ chính", "Đóng góp"], [
        ["[SV1 - MSSV]", "Thiết kế cơ sở dữ liệu, models, CRUD admin, báo cáo", "[...%]"],
        ["[SV2 - MSSV]", "Xác thực, đặt lịch, forms, templates, email", "[...%]"],
        ["[SV3 - MSSV]", "AI, API mạng, kiểm thử, Docker và triển khai", "[...%]"],
    ], [4.0, 9.8, 3.2])

    doc.add_heading("2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", level=1)
    doc.add_heading("2.1. Chức năng của hệ thống", level=2)
    paragraph(doc, "Các chức năng được phân tách theo vai trò để bảo đảm thao tác đúng quyền hạn. Khách có thể xem nội dung công khai, tra cứu dịch vụ và sử dụng chatbot. Thành viên có thể tạo lịch hẹn; quản trị viên được cấp thêm quyền quản lý dữ liệu và thống kê.")
    bullet(doc, "Khách: trang chủ, dịch vụ, đội ngũ, tin tức, liên hệ, chatbot AI.")
    bullet(doc, "Thành viên: đăng ký, đăng nhập, đặt lịch, xem lịch hẹn cá nhân.")
    bullet(doc, "Quản trị viên: dashboard, CRUD dịch vụ, duyệt/cập nhật trạng thái lịch hẹn, xem khách hàng và phản hồi.")
    bullet(doc, "Chức năng bắt buộc: đăng nhập/đăng ký, CRUD, tìm kiếm và phân trang.")
    caption(doc, "Hình 2.1. Sơ đồ Use Case (bổ sung hình vẽ: Khách, Thành viên và Quản trị viên).")

    doc.add_heading("2.2. Thiết kế cơ sở dữ liệu", level=2)
    paragraph(doc, "Cơ sở dữ liệu MySQL có tên nali_dental. Mô hình dữ liệu dùng các bảng phẳng để tương thích đồng thời với giao diện Flask và PHP. Bảng patients lưu khách hàng; users lưu nhân sự; products lưu dịch vụ; appointments lưu lịch; feedback lưu ý kiến khách hàng.")
    add_table(doc, "Bảng 2.1. Mô tả các bảng dữ liệu", ["Model/Bảng", "Cột chính", "Khóa và quan hệ"], [
        ["Patient / patients", "id, full_name, email, password, phone, gender", "id là PK; email duy nhất; một khách có nhiều lịch hẹn"],
        ["Staff / users", "id, username, password, full_name, role, specialty", "id là PK; role gồm admin, doctor, receptionist"],
        ["Product / products", "id, name, description, price, image, target_group, duration", "id là PK; danh mục dịch vụ nha khoa"],
        ["Appointment / appointments", "id, user_id, product_ids, customer_name, appointment_date, appointment_time, status", "id là PK; user_id tham chiếu khách hàng theo nghiệp vụ"],
        ["Feedback / feedback", "id, name, phone, email, rating, type, message", "id là PK; lưu phản hồi và mức hài lòng"],
    ], [3.4, 7.2, 6.4])
    paragraph(doc, "Quan hệ chính là Patient (1) - (n) Appointment. Một lịch có thể chứa danh sách dịch vụ qua product_ids; cách lưu này phù hợp schema dùng chung của dự án. Nhân sự quản trị và cập nhật trạng thái lịch hẹn thông qua giao diện admin.")
    caption(doc, "Hình 2.2. Sơ đồ ERD đề xuất (bổ sung hình vẽ quan hệ patients - appointments - products).")

    doc.add_heading("2.3. Thiết kế giao diện", level=2)
    paragraph(doc, "Ứng dụng Flask tổ chức giao diện bằng template kế thừa. base.html chứa header, footer, khu vực chatbot và các tài nguyên dùng chung; các trang con thuộc main, auth, booking và admin kế thừa khung này. Thẻ dịch vụ được đóng gói thành macro Jinja2 để dùng lại ở danh sách và trang liên quan.")
    paragraph(doc, "Luồng chính là: Trang chủ → danh sách dịch vụ (tìm kiếm, phân trang) → chi tiết dịch vụ → đăng nhập nếu cần → đặt lịch → lịch hẹn của tôi. Giao diện responsive, có dark mode, biểu đồ dashboard và widget chatbot. Hình ảnh dịch vụ được chuẩn hóa theo nhóm khách hàng để bảo đảm nhận diện thị giác nhất quán.")
    caption(doc, "Hình 2.3. Wireframe/luồng màn hình chính (bổ sung ảnh trang chủ, dịch vụ và đặt lịch).")

    doc.add_heading("3. XÂY DỰNG ỨNG DỤNG VỚI FLASK", level=1)
    doc.add_heading("3.1. Môi trường và cấu trúc dự án", level=2)
    paragraph(doc, "Dự án Flask được tổ chức theo Large Application Structure. Cấu hình được tách khỏi mã nguồn bằng biến môi trường; các extension được khởi tạo tập trung; route được chia thành Blueprint theo nghiệp vụ. Cách tổ chức này giúp mã nguồn dễ kiểm thử và mở rộng.")
    code_block(doc, "python -m venv .venv\n.venv\\Scripts\\activate\npip install -r requirements.txt\npython run.py")
    code_block(doc, "flask_app/\n├── run.py, config.py, requirements.txt\n└── app/\n    ├── __init__.py, extensions.py, models.py, forms.py\n    ├── auth.py, main.py, booking.py, admin.py, api.py\n    ├── templates/\n    └── static/")

    doc.add_heading("3.2. Routes và xử lý nghiệp vụ", level=2)
    paragraph(doc, "Các Blueprint gồm main_bp, auth_bp, booking_bp, admin_bp và api_bp. Route đăng nhập nhận email hoặc username, kiểm tra mật khẩu bcrypt rồi dùng Flask-Login để lưu phiên. Decorator phân quyền ngăn người dùng thường truy cập vùng /admin. Khi đặt lịch, hệ thống kiểm tra dữ liệu bắt buộc, ngày không thuộc quá khứ và giờ nằm trong khoảng 08:00-20:00.")
    code_block(doc, "@auth_bp.route(\"/dang-nhap\", methods=[\"GET\", \"POST\"])\ndef login():\n    form = LoginForm()\n    if form.validate_on_submit():\n        user = Patient.query.filter_by(email=form.email.data).first() \\\n               or Staff.query.filter_by(username=form.email.data).first()\n        if user and user.check_password(form.password.data):\n            login_user(user, remember=form.remember.data)\n            return redirect(url_for(\"main.index\"))")

    doc.add_heading("3.3. Templates với Jinja2", level=2)
    paragraph(doc, "Jinja2 được dùng để kế thừa giao diện, truyền dữ liệu từ controller và hiển thị có điều kiện theo quyền. Macro service_card giúp giảm lặp mã HTML. Danh sách dịch vụ nhận tham số q và page, sau đó render pagination; người dùng có thể tìm nhanh dịch vụ theo tên hoặc mô tả.")
    code_block(doc, "{% for p in services %}{{ service_card(p) }}{% endfor %}\n{% for n in pagination.iter_pages() %}\n  <a href=\"{{ url_for('main.services', q=q, page=n) }}\">{{ n }}</a>\n{% endfor %}")

    doc.add_heading("3.4. Web Forms với Flask-WTF", level=2)
    paragraph(doc, "Các lớp LoginForm, RegisterForm, AppointmentForm, ProductForm và FeedbackForm kế thừa FlaskForm. Hệ thống dùng DataRequired, Email, Length, Regexp, NumberRange và validator tự định nghĩa. CSRF được bật mặc định cho các form thay đổi dữ liệu; riêng form tìm kiếm sử dụng GET nên tắt CSRF.")
    code_block(doc, "class RegisterForm(FlaskForm):\n    email = StringField(\"Email\", validators=[DataRequired(), Email()])\n    phone = StringField(\"SĐT\", validators=[DataRequired(), Regexp(r\"^0\\d{9}$\")])\n    password = PasswordField(\"Mật khẩu\", validators=[DataRequired(), Length(6, 100)])")

    doc.add_heading("3.5. Cơ sở dữ liệu với Flask-SQLAlchemy", level=2)
    paragraph(doc, "Các model SQLAlchemy ánh xạ trực tiếp đến MySQL. Password tương thích hash bcrypt do PHP và Python tạo ra. Các thao tác thêm, sửa, xóa dùng session.add(), session.commit() và session.delete(); danh sách dịch vụ dùng query kết hợp filter và paginate để tránh tải toàn bộ dữ liệu trong một trang.")
    code_block(doc, "query = Product.query.filter(Product.name.like(f\"%{q}%\"))\npagination = query.paginate(page=page, per_page=6, error_out=False)\ndb.session.add(product)\ndb.session.commit()")

    doc.add_heading("3.6. Gửi email với Flask-Mail", level=2)
    paragraph(doc, "Sau khi ghi lịch hẹn, ứng dụng gọi helper send_email(). Hàm bọc Flask-Mail bằng try/except và ghi log nếu SMTP chưa cấu hình; vì vậy lỗi email không làm hỏng giao dịch đặt lịch. Thông tin SMTP được đọc qua biến môi trường để không đưa mật khẩu vào mã nguồn.")
    code_block(doc, "msg = Message(\"Xác nhận đặt lịch tại NALI\", recipients=[customer_email])\nmsg.body = f\"Mã lịch hẹn #{appointment.id}\"\nmail.send(msg)")

    doc.add_heading("3.7. Lập trình mạng và tích hợp API", level=2)
    paragraph(doc, "Yêu cầu CLO3 được thể hiện qua Blueprint api_bp. Endpoint /api/services trả danh sách dịch vụ ở JSON. Endpoint /api/weather gọi Open-Meteo bằng urllib.request, đọc JSON và trả nhiệt độ TP. Hồ Chí Minh. Endpoint /api/chat nhận JSON của giao diện, gửi POST đến FastAPI AI service bằng urllib, sau đó trả lại phản hồi cho trình duyệt.")
    code_block(doc, "with urllib.request.urlopen(url, timeout=8) as resp:\n    data = json.loads(resp.read().decode(\"utf-8\"))\nreturn jsonify({\"success\": True, \"temp\": data[\"current\"][\"temperature_2m\"]})")

    doc.add_heading("3.8. Chức năng mở rộng của nhóm", level=2)
    paragraph(doc, "Trợ lý AI được xây dựng thành FastAPI service độc lập. Retriever tạo kho tri thức từ dữ liệu phòng khám và dịch vụ, dùng TF-IDF khi không có API key hoặc embeddings khi có cấu hình phù hợp. Agent ưu tiên LLM Qwen2.5-3B đã finetune qua Ollama, sau đó Gemini và cuối cùng là agent offline theo luật. Agent offline vẫn hỗ trợ thu thập họ tên, số điện thoại, ngày, giờ và ghi lịch vào MySQL.")
    paragraph(doc, "Pipeline finetune gồm sinh dữ liệu SFT, train QLoRA, gộp adapter, xuất GGUF và nạp Ollama. Docker Compose khởi động bốn service: web PHP/Apache, MySQL, FastAPI AI và Ollama. Thiết kế nhiều tầng dự phòng giúp phần demo không bị gián đoạn khi model cục bộ hoặc dịch vụ ngoài chưa sẵn sàng.")

    doc.add_heading("4. KẾT QUẢ VÀ KIỂM THỬ", level=1)
    doc.add_heading("4.1. Kết quả giao diện", level=2)
    paragraph(doc, "Hệ thống đã hoàn thành các trang công khai, xác thực, đặt lịch, dashboard quản trị và chatbot. Bộ ảnh dịch vụ được chuẩn hóa về màu sắc xanh-trắng và bố cục để hiển thị đồng nhất trên lưới dịch vụ. Khi nộp báo cáo, nhóm cần chèn ảnh chụp giao diện thực tế tại các vị trí sau.")
    caption(doc, "Hình 4.1. Trang chủ NALI Dental.")
    caption(doc, "Hình 4.2. Danh sách dịch vụ có tìm kiếm, phân trang và ảnh đồng bộ.")
    caption(doc, "Hình 4.3. Đăng ký/đăng nhập và xác thực dữ liệu form.")
    caption(doc, "Hình 4.4. Form đặt lịch hẹn trực tuyến.")
    caption(doc, "Hình 4.5. Dashboard quản trị và quản lý lịch hẹn.")
    caption(doc, "Hình 4.6. Trợ lý AI tư vấn và đặt lịch theo hội thoại.")

    doc.add_heading("4.2. Kiểm thử chức năng", level=2)
    paragraph(doc, "Kiểm thử bao gồm unit/integration test Flask, bộ kiểm thử lõi AI offline và smoke test Docker. Kết quả tại thời điểm hoàn thiện: Flask đạt 10/10 test; AI đạt 23/23 test; các container web, MySQL, AI và Ollama đều khởi động thành công. Hai luồng đặt lịch qua form web và chatbot đã được kiểm tra ghi/xóa dữ liệu thử an toàn.")
    add_table(doc, "Bảng 4.1. Kết quả kiểm thử chức năng", ["TT", "Kịch bản", "Kết quả mong đợi", "Đánh giá"], [
        ["1", "Đăng ký tài khoản với email/SĐT hợp lệ", "Tạo khách hàng và cho phép đăng nhập", "Đạt"],
        ["2", "Đăng nhập sai mật khẩu", "Hiển thị lỗi, không tạo phiên", "Đạt"],
        ["3", "Tìm kiếm dịch vụ: implant", "Lọc đúng danh sách dịch vụ", "Đạt"],
        ["4", "Phân trang dịch vụ", "Điều hướng đúng trang dữ liệu", "Đạt"],
        ["5", "Đặt lịch bằng form", "Ghi MySQL, trả mã lịch hẹn", "Đạt"],
        ["6", "CRUD dịch vụ bởi admin", "Dữ liệu được thêm/sửa/xóa đúng quyền", "Đạt"],
        ["7", "GET /api/services", "Trả JSON danh sách dịch vụ", "Đạt"],
        ["8", "Hỏi AI giá tẩy trắng", "RAG trả thông tin dịch vụ phù hợp", "Đạt"],
        ["9", "Đặt lịch qua chatbot", "Agent kiểm tra và ghi lịch vào MySQL", "Đạt"],
        ["10", "Truy cập admin khi chưa đăng nhập", "Redirect/403 theo quyền", "Đạt"],
    ], [1.0, 6.5, 7.0, 2.0])

    doc.add_heading("5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    paragraph(doc, "Đề tài đã xây dựng thành công hệ thống quản lý phòng khám nha khoa NALI đáp ứng các yêu cầu cốt lõi của học phần: web Flask có cấu trúc module, template Jinja2, biểu mẫu Flask-WTF, dữ liệu SQLAlchemy/MySQL, xác thực, CRUD, tìm kiếm, phân trang, email và lập trình mạng với urllib/JSON. Nhóm cũng mở rộng hệ thống bằng chatbot AI, RAG, tool-calling đặt lịch và Docker Compose.")
    paragraph(doc, "Qua đề tài, nhóm củng cố kỹ năng thiết kế ứng dụng Flask theo Blueprint, xây dựng form an toàn, ánh xạ ORM, bảo vệ CSRF, xử lý API mạng, kiểm thử và đóng gói dịch vụ. Hạn chế hiện tại là SMTP cần cấu hình thực tế; LLM cục bộ chạy CPU còn chậm; và triển khai public cần hạ tầng có cơ sở dữ liệu bền vững.")
    bullet(doc, "Tích hợp thanh toán trực tuyến, xác nhận lịch qua OTP và lịch làm việc theo bác sĩ.")
    bullet(doc, "Bổ sung hồ sơ bệnh án điện tử, lịch sử điều trị và phân quyền chi tiết hơn.")
    bullet(doc, "Triển khai production với HTTPS, backup MySQL, giám sát log và tối ưu LLM bằng GPU/quantization.")

    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        "[1] M. Grinberg, Flask Web Development: Developing Web Applications with Python, 2nd ed. Sebastopol, CA: O'Reilly Media, 2018.",
        "[2] Pallets Projects, “Flask Documentation.” https://flask.palletsprojects.com/. [Truy cập: 23/07/2026].",
        "[3] Pallets Projects, “Flask-SQLAlchemy Documentation.” https://flask-sqlalchemy.palletsprojects.com/. [Truy cập: 23/07/2026].",
        "[4] Pallets Projects, “Flask-WTF Documentation.” https://flask-wtf.readthedocs.io/. [Truy cập: 23/07/2026].",
        "[5] Qwen Team, “Qwen2.5 Technical Report,” 2024. https://arxiv.org/abs/2412.15115. [Truy cập: 23/07/2026].",
        "[6] FastAPI, “FastAPI Documentation.” https://fastapi.tiangolo.com/. [Truy cập: 23/07/2026].",
        "[7] Docker, “Docker Compose Documentation.” https://docs.docker.com/compose/. [Truy cập: 23/07/2026].",
        "[8] Open-Meteo, “Weather Forecast API.” https://open-meteo.com/. [Truy cập: 23/07/2026].",
    ]
    for ref in refs:
        paragraph(doc, ref, first_line=False)

    doc.add_heading("PHỤ LỤC", level=1)
    doc.add_heading("Phụ lục A. Mã nguồn và hướng dẫn chạy", level=2)
    paragraph(doc, "Mã nguồn: https://github.com/MinhNhat-2504/NaLi-Dental-Clinic-Web. Để chạy bản Flask: vào thư mục flask_app, tạo môi trường ảo, cài requirements.txt và chạy python run.py. Tài khoản demo: admin/admin123 (quản trị), lananh@gmail.com/password123 (khách).")
    doc.add_heading("Phụ lục B. Dữ liệu mẫu và kiểm thử", level=2)
    paragraph(doc, "CSDL được tạo từ setup_database.php hoặc các lệnh init-db/seed-db của Flask. Bộ kiểm thử Flask chạy bằng pytest -q. Bộ kiểm thử AI chạy trong thư mục ai_service bằng python test_agent.py. Docker Compose được khởi động bằng docker compose up -d --build.")
    doc.add_heading("Phụ lục C. Hướng dẫn AI finetune", level=2)
    paragraph(doc, "Pipeline QLoRA và xuất GGUF được mô tả tại ai_service/finetune/README.md. Khi có file GGUF, nạp model vào Ollama bằng lệnh ollama create nali-dental -f finetune/Modelfile; AI service sẽ ưu tiên model này theo cấu hình LLM_BACKEND=local.")


def main():
    if not REPORT.exists():
        raise FileNotFoundError(REPORT)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)
    doc = Document(REPORT)
    setup_styles(doc)
    remove_template_body_after_toc(doc)
    add_content(doc)
    doc.core_properties.title = "Báo cáo đồ án môn học - NALI Dental"
    doc.core_properties.subject = "Lập trình Python nâng cao"
    doc.core_properties.author = "Nhóm thực hiện"
    doc.save(REPORT)
    print(REPORT)
    print(BACKUP)


if __name__ == "__main__":
    main()
